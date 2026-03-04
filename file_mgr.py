""" 
    Process files 
    includes saving and opening files.
    Handles conversion of data to Word Format.
"""
from tkinter import ttk, messagebox, filedialog
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import assembly

class FileManager:
    def __init__(self):
        # the dataframe is follows the notebook tabs.
        self.current_file = None    # to track the currently opened file.
        self.word_file = None       # to track the Word file

    # Create a static structure of the editor data dictionary.
    @staticmethod
    def default_editor_data():
        return {
            'Meeting Info' : {
                'Meeting Date' : '',
                'Start Time' : '',
            },
            'Opening Ceremony' : {
                'Intentions' : '',
                'Prayer' : '',
                'Leader' : '',
                'Pledge' : '',
            },
            'officers' : {
                "Faithful Navigator": {
                "name": "",
                "attendance": "",
                },
                "Faithful Friar": {
                    "name": "",
                    "attendance": "",
                },
                "Faithful Admiral": {
                    "name": "",
                    "attendance": "",
                },
                "Faithful Captain": {
                    "name": "",
                    "attendance": "",
                },
                "Faithful Pilot": {
                    "name": "",
                    "attendance": "",
                },
                "Faithful Comptroller": {
                    "name": "",
                    "attendance": "",
                },
                "Faithful Scribe": {
                    "name": "",
                    "attendance": "",
                },
                "Faithful Purser": {
                    "name": "",
                    "attendance": "",
                },
                "Faithful Inner Sentinel": {
                    "name": "",
                    "attendance": "",
                },
                "Faithful Outer Sentinel": {
                    "name": "",
                    "attendance": "",
                },
                "Faithful Trustee (1 Yr)": {
                    "name": "",
                    "attendance": "",
                },
                "Faithful Trustee (2 Yr)": {
                    "name": "",
                    "attendance": "",
                },
                "Faithful Trustee (3 Yr)": {
                    "name": "",
                    "attendance": "",
                },
            },
            "attendees": '',
            'Minutes' : {
                'Corrections' : '',
                'Motion to Approve' : '',
                'Seconded by' : '',
                'Approval' : '',
            },
            'Reports' : {
                'Friar' : '',
                'Bills' : '',
                'Comptroller' : '',
                'Purser' : '',
                'Standing Committees' : '',
                'Applications' : '',
                'Trustees' : '',
            },
            'Business' : {
                'Unfinished Business' : '',
                'New Business' : '',
            },
            'Council Reports' : {
                'Riverside' : '',
                'St Thomas' : '',
                'Sacred Heart' : '',
                'OLPH' : '',
            },
            'Closing Ceremony' : {
                'Closing Prayer' : '',
                'Intentions' : '',
                'Leader' : '',
                'Adjurned At' : '',
                'Next Officers Meeting': '',
                'Next Business Meeting' : '',
            },
            'Financials' : {
                'General': {
                    'Start Balance' : '',
                    'Receipts' : '',
                    'Deposits' : '',
                    'Disbursements' : '',
                    'End Balance': '',
                },
                'Chalice' : {
                    'Start Balance' : '',
                    'Receipts' : '',
                    'Deposits' : '',
                    'Disbursements' : '',
                    'End Balance': '',
                },
                'Flag' : {
                    'Start Balance' : '',
                    'Receipts' : '',
                    'Deposits' : '',
                    'Disbursements' : '',
                    'End Balance': '',
                },
            'Withdrawals' : '',
            'Deposits' : '',
            },
        }

    def new_json(self):
        """
            This function clears the dictionary editor_data and creates a new file.
        """
        
        
        print("The New Meeting Function to be developed.")

    def save_file_as_json(self, editor_data):
        """ This function saves the file under a new name."""
        # Open file dialog to choose location and name
        file_path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")],
            title="Save Meeting As"
        )
        if file_path:
            try:
                with open(file_path, 'w') as file_to_save:
                    json.dump(editor_data, file_to_save, indent=4)
                self.current_file = file_path
                print(f"Saved to: {file_path}")
                return True
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save file: {e}")
                print(f"Error saving: {e}")
                return False
        return False

    def save_to_file(self, editor_data):
        """" Saves a file. """
        if self.current_file:
            try:
                with open(self.current_file, 'w') as file_to_save:
                    json.dump(editor_data, file_to_save, indent=4)
                messagebox.showinfo("Success", f"The file <{self.current_file}> was saved successfully!")
                print(f"The file <{self.current_file}> was saved successfully!")
                return True
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save file: {e}")                
                print(f"Error saving {e}")
                return False
        else: 
            return self.save_file_as_json(editor_data)

    def open_json(self):
        """Load the working data from a user selected JSON file."""
        file_path = filedialog.askopenfilename(
            defaultextension='.json',
            filetypes=[("JSON files", '*.json'), ('All files', '*.*')],
            title="Open Meeting File"
        )

        if file_path:   # User didn't cancel
            try:
                with open(file_path, 'r') as f:
                    loaded_data = json.load(f)

                    # Update the editor_data with loaded data
                    self.current_file = file_path   # remember the file for later use

                    print(f"Loaded from: {file_path}")
                    messagebox.showinfo("Success", f"Meeting loaded data from: {file_path}")
                    return loaded_data
            except Exception as e:
                messagebox.showerror("Error", f"Failed to load file: {file_path}")
                print(f"Error loading: {e}")
                return None
        return None # if the user cancels the open dialog

    def convert_text_to_word(self, doc, paragraph_text):
        """ 
        Converts a string with '\n' into real paragrph breaks in a Word Document.
        Input: the Word dataframe, the text list
        Output: A bulletted sentence to the MS Word document
        """

        # Split the text into paragraphs based on '\n'
        paragraphs = paragraph_text.split('\n')
        
        for para in paragraphs:
            next_line = doc.add_paragraph(para.strip())
            next_line.style = 'List Bullet'

    def convert_to_word(self, editor_data):
        """ 
        This function converts a JSON into a MS Word document. 
        Import: the editor data dictionary. 
        Output: Word File
        """
        self.word_file = filedialog.asksaveasfilename(
            title="Export to Word",
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
            initialfile=f"KofC_Minutes_{datetime.now().strftime('%Y%m%d')}.docx"
        )
        print(f"Saving minutes to {self.word_file}")
        assembly_data = assembly.AssemblyInfo()

        # return if the user hits cancel.
        if not self.word_file:
            return
        
        # prepare the minutes for publication
        try:
            self.save_to_file(editor_data)
            doc = Document()


            # TODO: refactor this section of the code
            # prepare the title block
            title = doc.add_heading(f"Knights of Columbus {assembly_data.assembly_info['Assembly Name']} {assembly_data.assembly_info['Assembly Number']}", 1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle = doc.add_heading("Business Meeting Minutes", 2)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # List the Meeting Info
            doc.add_heading("Opening Ceremony", 2)
            date_para = doc.add_paragraph(f"Date: {editor_data['Meeting Info']['Meeting Date']}")
            date_para.space_after = Inches(0.0)
            doc.add_paragraph(f"Time: {editor_data['Meeting Info']['Start Time']}")

            # Print the body of the minutes.
            self._print_opening_ceremony_data(editor_data['Opening Ceremony'], doc)
            self._print_roll_call(roll_call=editor_data['officers'], other_attendees=editor_data['attendees'], doc=doc)
            self._assembly_report_data(editor_data['Reports'], editor_data['Minutes'], doc)
            self._print_business_report(editor_data['Business'], doc)
            self._print_council_report(report_data=editor_data['Council Reports'], doc=doc)
            self._print_closing_ceremony(editor_data['Closing Ceremony'], doc)
            self._create_financial_report(editor_data, doc, assembly_data)

            # Save the Document
            doc.save(self.word_file)
            messagebox.showinfo("Success", f"The Minutes were exported to:\n{self.word_file}")

        except Exception as e:
            messagebox.showerror('Export Error', f"An error occurred: {str(e)}")
    
    def _print_opening_ceremony_data(self, opening_data, doc):
        """Print the Opening Ceremony data to the Word Document"""
        doc.add_heading("Opening Ceremony", 2)
        doc.add_paragraph("Prayer Intentions:").bold = True
        self.convert_text_to_word(doc, opening_data['Intentions'])
        doc.add_paragraph(f"Prayer: {opening_data['Prayer']}")
        doc.add_paragraph(f"Prayer leg by: {opening_data['Leader']}")
        doc.add_paragraph(f"Pledge of Alligence led by: {opening_data['Pledge']}")
    
    def _print_roll_call(self, roll_call, other_attendees, doc):
        """"Print the meeting attendance information."""
        # Roll Call
        doc.add_heading("Roll Call", 2)
        table = doc.add_table(rows=13, cols=3)

        table.cell(0, 0).text = "Faithful Navigator"
        table.cell(0, 1).text = roll_call['Faithful Navigator']['name']
        table.cell(0, 2).text = roll_call['Faithful Navigator']['attendance']
        table.cell(1, 0).text = "Faithful Friar"
        table.cell(1, 1).text = roll_call['Faithful Friar']['name']
        table.cell(1, 2).text = roll_call['Faithful Friar']['attendance']
        table.cell(2, 0).text = "Faithful Admiral"
        table.cell(2, 1).text = roll_call['Faithful Admiral']['name']
        table.cell(2, 2).text = roll_call['Faithful Admiral']['attendance']
        table.cell(3, 0).text = "Faithful Captain"
        table.cell(3, 1).text = roll_call['Faithful Captain']['name']
        table.cell(3, 2).text = roll_call['Faithful Captain']['attendance']
        table.cell(4, 0).text = "Faithful Pilot"
        table.cell(4, 1).text = roll_call['Faithful Pilot']['name']
        table.cell(4, 2).text = roll_call['Faithful Pilot']['attendance']
        table.cell(5, 0).text = "Faithful Comptroller"
        table.cell(5, 1).text = roll_call['Faithful Comptroller']['name']
        table.cell(5, 2).text = roll_call['Faithful Comptroller']['attendance']
        table.cell(6, 0).text = "Faithful Scribe"
        table.cell(6, 1).text = roll_call['Faithful Scribe']['name']
        table.cell(6, 2).text = roll_call['Faithful Scribe']['attendance']
        table.cell(7, 0).text = "Faithful Purser"
        table.cell(7, 1).text = roll_call['Faithful Purser']['name']
        table.cell(7, 2).text = roll_call['Faithful Purser']['attendance']
        table.cell(8, 0).text = "Faithful Inner Sentinel"
        table.cell(8, 1).text = roll_call['Faithful Inner Sentinel']['name']
        table.cell(8, 2).text = roll_call['Faithful Inner Sentinel']['attendance']
        table.cell(9, 0).text = "Faithful Outer Sentinel"
        table.cell(9, 1).text = roll_call['Faithful Outer Sentinel']['name']
        table.cell(9, 2).text = roll_call['Faithful Outer Sentinel']['attendance']
        table.cell(10, 0).text = "Faithful Trustee (1 Yr)"
        table.cell(10, 1).text = roll_call['Faithful Trustee (1 Yr)']['name']
        table.cell(10, 2).text = roll_call['Faithful Trustee (1 Yr)']['attendance']
        table.cell(11, 0).text = "Faithful Trustee (2 Yr)"
        table.cell(11, 1).text = roll_call['Faithful Trustee (2 Yr)']['name']
        table.cell(11, 2).text = roll_call['Faithful Trustee (2 Yr)']['attendance']
        table.cell(12, 0).text = "Faithful Trustee (3 Yr)"
        table.cell(12, 1).text = roll_call['Faithful Trustee (3 Yr)']['name']
        table.cell(12, 2).text = roll_call['Faithful Trustee (3 Yr)']['attendance']
        # Other Attendees
        if other_attendees != '':
            doc.add_paragraph("Members in attendance:")
            self.convert_text_to_word(doc, other_attendees)

    def _assembly_report_data(self, assembly_reports, minute_report, doc):
        """Print the Assembly Reports to the Word document."""
        # Start with the Faithful Friar's report
        doc.add_heading("Faithful Friar's Report", 2)
        self.convert_text_to_word(doc, assembly_reports['Friar'])

        # Handle the minutes of the last meeting.
        doc.add_heading("Reading of Minutes", 2)
        self.convert_text_to_word(doc, minute_report['Corrections'])
        doc.add_paragraph(f"Motion to approve by: {minute_report['Motion to Approve']}")
        doc.add_paragraph(f"Seconded by: {minute_report['Seconded by']}")
        doc.add_paragraph(f"Minutes approved: {minute_report['Approval']}")

        doc.add_heading("Bills and Communications", 2)
        self.convert_text_to_word(doc, assembly_reports['Bills'])

        doc.add_heading("Faithful Comptroller's Report", 2)
        self.convert_text_to_word(doc, assembly_reports['Comptroller'])

        doc.add_heading("Purser's Report", 2)
        self.convert_text_to_word(doc,assembly_reports['Purser'])

        doc.add_heading("Standing Committees", 2)
        doc.add_heading("Color Corps Commander", 3)
        self.convert_text_to_word(doc,assembly_reports['Standing Committees'])

        doc.add_heading("Reading of Applications", 2)
        self.convert_text_to_word(doc, assembly_reports['Applications'])

        doc.add_heading("Trustees Report", 2)
        self.convert_text_to_word(doc, assembly_reports['Trustees'])
    
    def _print_business_report(self, assy_business_data, doc):
        """Print the Assembly Business reports to the Word document."""
        doc.add_heading("Unfinished Business", 2)
        self.convert_text_to_word(doc, assy_business_data['Unfinished Business'])
        doc.add_heading("New Business", 2)
        self.convert_text_to_word(doc, assy_business_data['New Business'])
        

    def _print_council_report(self, report_data, doc):
        """Prints the Council reports to the Word Document."""
        doc.add_heading("Report of the Councils", 2)
        doc.add_heading("Riverside", 3)
        self.convert_text_to_word(doc,report_data['Riverside'])
        doc.add_heading("Saint Thomas", 3)
        self.convert_text_to_word(doc, report_data['St Thomas'])
        doc.add_heading("Sacred Heart", 3)
        self.convert_text_to_word(doc, report_data['Sacred Heart'])
        doc.add_heading("Our Lady of Perpetual Health", 3)
        self.convert_text_to_word(doc, report_data['OLPH'])

    def _print_closing_ceremony(self, closing_ceremony_data, doc):
        """ Print the Closing Ceremony to the Word Document."""
        doc.add_heading("Closing Ceremony", 2)
        doc.add_paragraph(f"Closing prayer led by: {closing_ceremony_data['Leader']}")
        doc.add_paragraph(f"Closing Prayer: {closing_ceremony_data['Closing Prayer']}")
        doc.add_paragraph("Prayer Intentions:")
        self.convert_text_to_word(doc, closing_ceremony_data['Intentions'])
        doc.add_paragraph(f"Next Officers meeting: {closing_ceremony_data['Next Officers Meeting']}")
        doc.add_paragraph(f"Next Business Meeting: {closing_ceremony_data['Next Business Meeting']}")

    def _create_financial_report(self, editor_data, doc, assembly_data):
        """Create the Financial Report section in the Wod Document."""
        # Start on a new page
        doc.add_page_break()
        doc.add_heading(f"Knights of Columbus {assembly_data.assembly_info['Assembly Name']} {assembly_data.assembly_info['Assembly Number']}", 2).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_heading("FAITHFUL PURSER'S MONTHLY FINANCIAL REPORT", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"As of {editor_data['Meeting Info']['Meeting Date']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_heading("General Fund Account", 3).alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._print_general_report(editor_data, doc)
        doc.add_heading("Chalice Special Fund Account", 3).alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._print_chalice_report(editor_data, doc)
        doc.add_heading("Flag Special Fund Account", 3).alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._print_flag_report(editor_data, doc)
        self._print_transactions(editor_data, doc)

    def _print_general_report(self, editor_data, doc):
        """Print the General Financial Report to the Word Document."""
        general_table = doc.add_table(rows=5, cols=2)
        general_table.cell(0, 0).text="Starting Monthly Balance"
        general_table.cell(0, 1).text=editor_data['Financials']['General']['Start Balance']
        general_table.cell(1, 0).text="Total Receipts"
        general_table.cell(1, 1).text=editor_data['Financials']['General']['Receipts']
        general_table.cell(2, 0).text="Funds Deposited (membership)"
        general_table.cell(2, 1).text=editor_data['Financials']['General']['Deposits']
        general_table.cell(3, 0).text="Less Total Disbursements"
        general_table.cell(3, 1).text=editor_data['Financials']['General']['Disbursements']
        general_table.cell(4, 0).text="Ending Balance"
        general_table.cell(4, 1).text=editor_data['Financials']['General']['End Balance']

    def _print_chalice_report(self, editor_data, doc):
        """Print the Chalice Fincalcial Report to the Word Document."""
        chalice_table = doc.add_table(rows=5, cols=2)
        chalice_table.cell(0, 0).text="Starting Monthly Balance"
        chalice_table.cell(0, 1).text=editor_data['Financials']['Chalice']['Start Balance']
        chalice_table.cell(1, 0).text="Total Receipts"
        chalice_table.cell(1, 1).text=editor_data['Financials']['Chalice']['Receipts']
        chalice_table.cell(2, 0).text="Funds Deposited"
        chalice_table.cell(2, 1).text=editor_data['Financials']['Chalice']['Deposits']
        chalice_table.cell(3, 0).text="Less Total Disbursements"
        chalice_table.cell(3, 1).text=editor_data['Financials']['Chalice']['Disbursements']
        chalice_table.cell(4, 0).text="Ending Balance"
        chalice_table.cell(4, 1).text=editor_data['Financials']['Chalice']['End Balance']

    def _print_flag_report(self, editor_data, doc):
        """Print the Flag Financial Report to the Word Document."""
        flag_table = doc.add_table(rows=5, cols=2)
        flag_table.cell(0, 0).text="Starting Monthly Balance"
        flag_table.cell(0, 1).text=editor_data['Financials']['Flag']['Start Balance']
        flag_table.cell(1, 0).text="Total Receipts"
        flag_table.cell(1, 1).text=editor_data['Financials']['Flag']['Receipts']
        flag_table.cell(2, 0).text="Funds Deposited"
        flag_table.cell(2, 1).text=editor_data['Financials']['Flag']['Deposits']
        flag_table.cell(3, 0).text="Less Total Disbursements"
        flag_table.cell(3, 1).text=editor_data['Financials']['Flag']['Disbursements']
        flag_table.cell(4, 0).text="Ending Balance"
        flag_table.cell(4, 1).text=editor_data['Financials']['Flag']['End Balance']

    def _print_transactions(self, editor_data, doc):
        """Print the list of transactions to the Word Document."""
        if editor_data['Financials']['Withdrawals'] != '':
            doc.add_heading("Withdrawal",3)
            self.convert_text_to_word(doc, editor_data['Financials']['Withdrawals'])
        if editor_data['Financials']['Deposits'] != '':
            doc.add_heading("Deposits", 3)
            self.convert_text_to_word(doc, editor_data['Financials']['Deposits'])
