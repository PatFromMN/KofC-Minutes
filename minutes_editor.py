import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os


class KofCMinutesApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Knights of Columbus 4th Degree - Meeting Minutes")
        self.root.geometry("1000x800")
        
        self.current_file = None
        self.data = self.initialize_data()
        
        # Create menu
        menubar = tk.Menu(root)
        root.config(menu=menubar)
        
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="New Meeting", command=self.new_meeting)
        file_menu.add_command(label="Open Working File", command=self.open_file)
        file_menu.add_command(label="Save Working File", command=self.save_file)
        file_menu.add_command(label="Save As...", command=self.save_file_as)
        file_menu.add_separator()
        file_menu.add_command(label="Export to Word", command=self.export_to_word)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=root.quit)
        
        # Create notebook for tabs
        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Create tabs
        self.create_meeting_info_tab()
        self.create_opening_ceremony_tab()
        self.create_roll_call_tab()
        self.create_reports_tab()
        self.create_business_tab()
        self.create_closing_tab()
        
    def initialize_data(self):
        """Initialize the data structure"""
        return {
            "meeting_data": {
                "date": datetime.now().strftime("%Y-%m-%d"),
                "time": "",
            },
            "opening_ceremony": {
                "prayer_intentions": "",
                "opening_prayer": "",
                "led_by": "",
                "pledge_of_allegiance": "",
            },
            "roll_call": {
                "Faithful Navigator": {
                    "name": "Ruben J. Rosas, PFN",
                    "attendance": "Present",
                },
                "Faithful Friar": {
                    "name": "Father Paul L. Schmidt",
                    "attendance": "Present",
                },
                "Faithful Admiral": {
                    "name": "Kurt Forster, PFN",
                    "attendance": "Present",
                },
                "Faithful Captain": {
                    "name": "Richard Aceves",
                    "attendance": "Present",
                },
                "Faithful Pilot": {
                    "name": "Paul Avalos",
                    "attendance": "Present",
                },
                "Faithful Comptroller": {
                    "name": "Thomas Walsh, PFN",
                    "attendance": "Present",
                },
                "Faithful Scribe": {
                    "name": "Patrick Norris",
                    "attendance": "Present",
                },
                "Faithful Purser": {
                    "name": "Lincoln Mena",
                    "attendance": "Present",
                },
                "Faithful Inner Sentinel": {
                    "name": "Lonnie Kobus",
                    "attendance": "Present",
                },
                "Faithful Outer Sentinel": {
                    "name": "Romulus Fundeanu",
                    "attendance": "Present",
                },
                "Faithful Trustee (1 Yr)": {
                    "name": "Vacant",
                    "attendance": "Not Applicable",
                },
                "Faithful Trustee (2-Yr)": {
                    "name": "Lester Pastuszyn",
                    "attendance": "Present",
                },
                "Faithful Trustee (3-Yr)": {
                    "name": "Rene Luna",
                    "attendance": "Present",
                },
            },
            "meeting_notes": {
                "Faithful Friar": "",
                "Reading of Minutes": {
                    "corrections": "",
                    "motion_to_approve": "",
                    "seconded_by": "",
                    "approved": "Yes",
                },
                "Bills and Communications": "",
                "Faithful Comptroller": "",
                "Purser's Report": "",
                "Standing Committees": {
                    "Color Corps": ""
                },
                "Reading of Applications": "",
                "Unfinished Business": "",
                "New Business": "",
                "Trustee's Report": "",
                "Report of the Councils": {
                    "Riverside": "",
                    "Saint Thomas": "",
                    "Sacred Heart": "",
                    "Our Lady of Perpetual Help": "",
                },
            },
            "closing_notes": {
                "time": "",
                "Next Officers Meeting": "",
                "Next Business Meeting": "",
            },
        }
    
    def create_meeting_info_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Meeting Info")
        
        content = ttk.Frame(frame, padding="20")
        content.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(content, text="Date:", font=("Arial", 10, "bold")).grid(row=0, column=0, sticky=tk.W, pady=5)
        self.date_entry = ttk.Entry(content, width=30)
        self.date_entry.grid(row=0, column=1, sticky=tk.W, pady=5)
        
        ttk.Label(content, text="Starting Time:", font=("Arial", 10, "bold")).grid(row=1, column=0, sticky=tk.W, pady=5)
        self.time_entry = ttk.Entry(content, width=30)
        self.time_entry.grid(row=1, column=1, sticky=tk.W, pady=5)
        
    def create_opening_ceremony_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Opening Ceremony")
        
        content = ttk.Frame(frame, padding="20")
        content.pack(fill=tk.BOTH, expand=True)
        
        row = 0
        ttk.Label(content, text="Prayer Intentions:", font=("Arial", 10, "bold")).grid(row=row, column=0, sticky=tk.NW, pady=5)
        self.prayer_intentions = scrolledtext.ScrolledText(content, width=60, height=4)
        self.prayer_intentions.grid(row=row, column=1, pady=5, sticky=(tk.W, tk.E))
        
        row += 1
        ttk.Label(content, text="Opening Prayer:", font=("Arial", 10, "bold")).grid(row=row, column=0, sticky=tk.W, pady=5)
        self.opening_prayer = ttk.Entry(content, width=60)
        self.opening_prayer.grid(row=row, column=1, pady=5, sticky=(tk.W, tk.E))
        
        row += 1
        ttk.Label(content, text="Led By:", font=("Arial", 10, "bold")).grid(row=row, column=0, sticky=tk.W, pady=5)
        self.opening_led_by = ttk.Entry(content, width=60)
        self.opening_led_by.grid(row=row, column=1, pady=5, sticky=(tk.W, tk.E))
        
        row += 1
        ttk.Label(content, text="Pledge of Allegiance Led By:", font=("Arial", 10, "bold")).grid(row=row, column=0, sticky=tk.NW, pady=5)
        self.pledge_led_by = ttk.Entry(content, width=60)
        self.pledge_led_by.grid(row=row, column=1, pady=5, sticky=(tk.W, tk.E))
        
        content.columnconfigure(1, weight=1)
        
    def create_roll_call_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Roll Call")
        
        # Create canvas with scrollbar
        canvas = tk.Canvas(frame)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        content = ttk.Frame(scrollable_frame, padding="20")
        content.pack(fill=tk.BOTH, expand=True)
        
        # Header
        ttk.Label(content, text="Office", font=("Arial", 10, "bold")).grid(row=0, column=0, padx=5, pady=5)
        ttk.Label(content, text="Name", font=("Arial", 10, "bold")).grid(row=0, column=1, padx=5, pady=5)
        ttk.Label(content, text="Attendance", font=("Arial", 10, "bold")).grid(row=0, column=2, columnspan=3, padx=5, pady=5)
        
        self.attendance_vars = {}
        row = 1
        
        for office, info in self.data["roll_call"].items():
            ttk.Label(content, text=office).grid(row=row, column=0, sticky=tk.W, padx=5, pady=2)
            ttk.Label(content, text=info["name"]).grid(row=row, column=1, sticky=tk.W, padx=5, pady=2)
            
            var = tk.StringVar(value=info["attendance"])
            self.attendance_vars[office] = var
            
            ttk.Radiobutton(content, text="Present", variable=var, value="Present").grid(row=row, column=2, padx=2)
            ttk.Radiobutton(content, text="Absent", variable=var, value="Absent").grid(row=row, column=3, padx=2)
            ttk.Radiobutton(content, text="Excused", variable=var, value="Excused").grid(row=row, column=4, padx=2)
            
            if "Vacant" in info["name"]:
                ttk.Radiobutton(content, text="N/A", variable=var, value="Not Applicable").grid(row=row, column=5, padx=2)
            
            row += 1
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
    def create_reports_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Reports")
        
        # Create canvas with scrollbar
        canvas = tk.Canvas(frame)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        content = ttk.Frame(scrollable_frame, padding="20")
        content.pack(fill=tk.BOTH, expand=True)
        
        row = 0
        
        # Faithful Friar
        ttk.Label(content, text="Faithful Friar:", font=("Arial", 10, "bold")).grid(row=row, column=0, sticky=tk.NW, pady=5)
        self.friar_report = scrolledtext.ScrolledText(content, width=70, height=4)
        self.friar_report.grid(row=row, column=1, pady=5, sticky=(tk.W, tk.E))
        row += 1
        
        # Reading of Minutes
        ttk.Label(content, text="Reading of Minutes", font=("Arial", 11, "bold")).grid(row=row, column=0, columnspan=2, sticky=tk.W, pady=(10, 5))
        row += 1
        
        ttk.Label(content, text="Corrections:").grid(row=row, column=0, sticky=tk.NW, pady=2, padx=(20, 0))
        self.minutes_corrections = scrolledtext.ScrolledText(content, width=70, height=2)
        self.minutes_corrections.grid(row=row, column=1, pady=2, sticky=(tk.W, tk.E))
        row += 1
        
        ttk.Label(content, text="Motion to Approve:").grid(row=row, column=0, sticky=tk.W, pady=2, padx=(20, 0))
        self.minutes_motion = ttk.Entry(content, width=70)
        self.minutes_motion.grid(row=row, column=1, pady=2, sticky=(tk.W, tk.E))
        row += 1
        
        ttk.Label(content, text="Seconded By:").grid(row=row, column=0, sticky=tk.W, pady=2, padx=(20, 0))
        self.minutes_seconded = ttk.Entry(content, width=70)
        self.minutes_seconded.grid(row=row, column=1, pady=2, sticky=(tk.W, tk.E))
        row += 1
        
        ttk.Label(content, text="Approved:").grid(row=row, column=0, sticky=tk.W, pady=2, padx=(20, 0))
        self.minutes_approved = ttk.Combobox(content, values=["Yes", "No", "Tabled"], width=20)
        self.minutes_approved.grid(row=row, column=1, pady=2, sticky=tk.W)
        row += 1
        
        # Bills and Communications
        ttk.Label(content, text="Bills and Communications:", font=("Arial", 10, "bold")).grid(row=row, column=0, sticky=tk.NW, pady=5)
        self.bills_comm = scrolledtext.ScrolledText(content, width=70, height=3)
        self.bills_comm.grid(row=row, column=1, pady=5, sticky=(tk.W, tk.E))
        row += 1
        
        # Comptroller
        ttk.Label(content, text="Faithful Comptroller:", font=("Arial", 10, "bold")).grid(row=row, column=0, sticky=tk.NW, pady=5)
        self.comptroller_report = scrolledtext.ScrolledText(content, width=70, height=3)
        self.comptroller_report.grid(row=row, column=1, pady=5, sticky=(tk.W, tk.E))
        row += 1
        
        # Purser
        ttk.Label(content, text="Purser's Report:", font=("Arial", 10, "bold")).grid(row=row, column=0, sticky=tk.NW, pady=5)
        self.purser_report = scrolledtext.ScrolledText(content, width=70, height=3)
        self.purser_report.grid(row=row, column=1, pady=5, sticky=(tk.W, tk.E))
        row += 1
        
        # Standing Committees
        ttk.Label(content, text="Standing Committees", font=("Arial", 11, "bold")).grid(row=row, column=0, columnspan=2, sticky=tk.W, pady=(10, 5))
        row += 1
        
        ttk.Label(content, text="Color Corps:").grid(row=row, column=0, sticky=tk.NW, pady=2, padx=(20, 0))
        self.color_corps = scrolledtext.ScrolledText(content, width=70, height=3)
        self.color_corps.grid(row=row, column=1, pady=2, sticky=(tk.W, tk.E))
        row += 1
        
        # Applications
        ttk.Label(content, text="Reading of Applications:", font=("Arial", 10, "bold")).grid(row=row, column=0, sticky=tk.NW, pady=5)
        self.applications = scrolledtext.ScrolledText(content, width=70, height=3)
        self.applications.grid(row=row, column=1, pady=5, sticky=(tk.W, tk.E))
        row += 1
        
        # Trustee's Report
        ttk.Label(content, text="Trustee's Report:", font=("Arial", 10, "bold")).grid(row=row, column=0, sticky=tk.NW, pady=5)
        self.trustee_report = scrolledtext.ScrolledText(content, width=70, height=3)
        self.trustee_report.grid(row=row, column=1, pady=5, sticky=(tk.W, tk.E))
        row += 1
        
        # Report of the Councils
        ttk.Label(content, text="Report of the Councils", font=("Arial", 11, "bold")).grid(row=row, column=0, columnspan=2, sticky=tk.W, pady=(10, 5))
        row += 1
        
        self.council_reports = {}
        for council in ["Riverside", "Saint Thomas", "Sacred Heart", "Our Lady of Perpetual Help"]:
            ttk.Label(content, text=f"{council}:").grid(row=row, column=0, sticky=tk.NW, pady=2, padx=(20, 0))
            text_widget = scrolledtext.ScrolledText(content, width=70, height=2)
            text_widget.grid(row=row, column=1, pady=2, sticky=(tk.W, tk.E))
            self.council_reports[council] = text_widget
            row += 1
        
        content.columnconfigure(1, weight=1)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
    def create_business_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Business")
        
        content = ttk.Frame(frame, padding="20")
        content.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(content, text="Unfinished Business:", font=("Arial", 10, "bold")).pack(anchor=tk.W, pady=(5, 2))
        self.unfinished_business = scrolledtext.ScrolledText(content, width=80, height=12)
        self.unfinished_business.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        ttk.Label(content, text="New Business:", font=("Arial", 10, "bold")).pack(anchor=tk.W, pady=(5, 2))
        self.new_business = scrolledtext.ScrolledText(content, width=80, height=12)
        self.new_business.pack(fill=tk.BOTH, expand=True)
        
    def create_closing_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Closing")
        
        content = ttk.Frame(frame, padding="20")
        content.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(content, text="Closing Time:", font=("Arial", 10, "bold")).grid(row=0, column=0, sticky=tk.W, pady=5)
        self.closing_time = ttk.Entry(content, width=30)
        self.closing_time.grid(row=0, column=1, sticky=tk.W, pady=5)
        
        ttk.Label(content, text="Next Officers Meeting:", font=("Arial", 10, "bold")).grid(row=1, column=0, sticky=tk.W, pady=5)
        self.next_officers = ttk.Entry(content, width=50)
        self.next_officers.grid(row=1, column=1, sticky=tk.W, pady=5)
        
        ttk.Label(content, text="Next Business Meeting:", font=("Arial", 10, "bold")).grid(row=2, column=0, sticky=tk.W, pady=5)
        self.next_business = ttk.Entry(content, width=50)
        self.next_business.grid(row=2, column=1, sticky=tk.W, pady=5)
        
    def load_data_to_gui(self):
        """Load data from self.data into GUI widgets"""
        # Meeting Info
        self.date_entry.delete(0, tk.END)
        self.date_entry.insert(0, self.data["meeting_data"]["date"])
        self.time_entry.delete(0, tk.END)
        self.time_entry.insert(0, self.data["meeting_data"]["time"])
        
        # Opening Ceremony
        self.prayer_intentions.delete(1.0, tk.END)
        self.prayer_intentions.insert(1.0, self.data["opening_ceremony"]["prayer_intentions"])
        self.opening_prayer.delete(0, tk.END)
        self.opening_prayer.insert(0, self.data["opening_ceremony"]["opening_prayer"])
        self.opening_led_by.delete(0, tk.END)
        self.opening_led_by.insert(0, self.data["opening_ceremony"]["led_by"])
        self.pledge_led_by.delete(0, tk.END)
        self.pledge_led_by.insert(0, self.data["opening_ceremony"]["pledge_of_allegiance"])
        
        # Roll Call
        for office, var in self.attendance_vars.items():
            var.set(self.data["roll_call"][office]["attendance"])
        
        # Reports
        self.friar_report.delete(1.0, tk.END)
        self.friar_report.insert(1.0, self.data["meeting_notes"]["Faithful Friar"])
        
        self.minutes_corrections.delete(1.0, tk.END)
        self.minutes_corrections.insert(1.0, self.data["meeting_notes"]["Reading of Minutes"]["corrections"])
        self.minutes_motion.delete(0, tk.END)
        self.minutes_motion.insert(0, self.data["meeting_notes"]["Reading of Minutes"]["motion_to_approve"])
        self.minutes_seconded.delete(0, tk.END)
        self.minutes_seconded.insert(0, self.data["meeting_notes"]["Reading of Minutes"]["seconded_by"])
        self.minutes_approved.set(self.data["meeting_notes"]["Reading of Minutes"]["approved"])
        
        self.bills_comm.delete(1.0, tk.END)
        self.bills_comm.insert(1.0, self.data["meeting_notes"]["Bills and Communications"])
        
        self.comptroller_report.delete(1.0, tk.END)
        self.comptroller_report.insert(1.0, self.data["meeting_notes"]["Faithful Comptroller"])
        
        self.purser_report.delete(1.0, tk.END)
        self.purser_report.insert(1.0, self.data["meeting_notes"]["Purser's Report"])
        
        self.color_corps.delete(1.0, tk.END)
        self.color_corps.insert(1.0, self.data["meeting_notes"]["Standing Committees"]["Color Corps"])
        
        self.applications.delete(1.0, tk.END)
        self.applications.insert(1.0, self.data["meeting_notes"]["Reading of Applications"])
        
        self.trustee_report.delete(1.0, tk.END)
        self.trustee_report.insert(1.0, self.data["meeting_notes"]["Trustee's Report"])
        
        for council, widget in self.council_reports.items():
            widget.delete(1.0, tk.END)
            widget.insert(1.0, self.data["meeting_notes"]["Report of the Councils"][council])
        
        # Business
        self.unfinished_business.delete(1.0, tk.END)
        self.unfinished_business.insert(1.0, self.data["meeting_notes"]["Unfinished Business"])
        
        self.new_business.delete(1.0, tk.END)
        self.new_business.insert(1.0, self.data["meeting_notes"]["New Business"])
        
        # Closing
        self.closing_time.delete(0, tk.END)
        self.closing_time.insert(0, self.data["closing_notes"]["time"])
        self.next_officers.delete(0, tk.END)
        self.next_officers.insert(0, self.data["closing_notes"]["Next Officers Meeting"])
        self.next_business.delete(0, tk.END)
        self.next_business.insert(0, self.data["closing_notes"]["Next Business Meeting"])
        
    def save_gui_to_data(self):
        """Save GUI data back to self.data"""
        # Meeting Info
        self.data["meeting_data"]["date"] = self.date_entry.get()
        self.data["meeting_data"]["time"] = self.time_entry.get()
        
        # Opening Ceremony
        self.data["opening_ceremony"]["prayer_intentions"] = self.prayer_intentions.get(1.0, tk.END).strip()
        self.data["opening_ceremony"]["opening_prayer"] = self.opening_prayer.get()
        self.data["opening_ceremony"]["led_by"] = self.opening_led_by.get()
        self.data["opening_ceremony"]["pledge_of_allegiance"] = self.pledge_led_by.get()
        
        # Roll Call
        for office, var in self.attendance_vars.items():
            self.data["roll_call"][office]["attendance"] = var.get()
        
        # Reports
        self.data["meeting_notes"]["Faithful Friar"] = self.friar_report.get(1.0, tk.END).strip()
        self.data["meeting_notes"]["Reading of Minutes"]["corrections"] = self.minutes_corrections.get(1.0, tk.END).strip()
        self.data["meeting_notes"]["Reading of Minutes"]["motion_to_approve"] = self.minutes_motion.get()
        self.data["meeting_notes"]["Reading of Minutes"]["seconded_by"] = self.minutes_seconded.get()
        self.data["meeting_notes"]["Reading of Minutes"]["approved"] = self.minutes_approved.get()
        self.data["meeting_notes"]["Bills and Communications"] = self.bills_comm.get(1.0, tk.END).strip()
        self.data["meeting_notes"]["Faithful Comptroller"] = self.comptroller_report.get(1.0, tk.END).strip()
        self.data["meeting_notes"]["Purser's Report"] = self.purser_report.get(1.0, tk.END).strip()
        self.data["meeting_notes"]["Standing Committees"]["Color Corps"] = self.color_corps.get(1.0, tk.END).strip()
        self.data["meeting_notes"]["Reading of Applications"] = self.applications.get(1.0, tk.END).strip()
        self.data["meeting_notes"]["Trustee's Report"] = self.trustee_report.get(1.0, tk.END).strip()
        
        for council, widget in self.council_reports.items():
            self.data["meeting_notes"]["Report of the Councils"][council] = widget.get(1.0, tk.END).strip()
        
        # Business
        self.data["meeting_notes"]["Unfinished Business"] = self.unfinished_business.get(1.0, tk.END).strip()
        self.data["meeting_notes"]["New Business"] = self.new_business.get(1.0, tk.END).strip()
        
        # Closing
        self.data["closing_notes"]["time"] = self.closing_time.get()
        self.data["closing_notes"]["Next Officers Meeting"] = self.next_officers.get()
        self.data["closing_notes"]["Next Business Meeting"] = self.next_business.get()
        
    def new_meeting(self):
        if messagebox.askyesno("New Meeting", "Start a new meeting? Any unsaved changes will be lost."):
            self.data = self.initialize_data()
            self.current_file = None
            self.load_data_to_gui()
            self.root.title("Knights of Columbus 4th Degree - Meeting Minutes [New]")
            
    def open_file(self):
        filename = filedialog.askopenfilename(
            title="Open Working File",
            filetypes=[("JSON Files", "*.json"), ("All Files", "*.*")]
        )
        if filename:
            try:
                with open(filename, 'r') as f:
                    self.data = json.load(f)
                self.current_file = filename
                self.load_data_to_gui()
                self.root.title(f"Knights of Columbus 4th Degree - Meeting Minutes [{os.path.basename(filename)}]")
                messagebox.showinfo("Success", "File loaded successfully!")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to open file:\n{str(e)}")
                
    def save_file(self):
        if self.current_file:
            self.save_to_file(self.current_file)
        else:
            self.save_file_as()
            
    def save_file_as(self):
        filename = filedialog.asksaveasfilename(
            title="Save Working File",
            defaultextension=".json",
            filetypes=[("JSON Files", "*.json"), ("All Files", "*.*")],
            initialfile=f"KofC_Minutes_{datetime.now().strftime('%Y%m%d')}.json"
        )
        if filename:
            self.save_to_file(filename)
            
    def save_to_file(self, filename):
        try:
            self.save_gui_to_data()
            with open(filename, 'w') as f:
                json.dump(self.data, f, indent=4)
            self.current_file = filename
            self.root.title(f"Knights of Columbus 4th Degree - Meeting Minutes [{os.path.basename(filename)}]")
            messagebox.showinfo("Success", "File saved successfully!")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save file:\n{str(e)}")
            
    def export_to_word(self):
        filename = filedialog.asksaveasfilename(
            title="Export to Word",
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
            initialfile=f"KofC_Minutes_{datetime.now().strftime('%Y%m%d')}.docx"
        )
        
        if not filename:
            return
            
        try:
            self.save_gui_to_data()
            doc = Document()
            
            # Title
            title = doc.add_heading("Knights of Columbus 4th Degree", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle = doc.add_heading("Meeting Minutes", level=2)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Meeting Info
            doc.add_paragraph(f"Date: {self.data['meeting_data']['date']}")
            doc.add_paragraph(f"Time: {self.data['meeting_data']['time']}")
            doc.add_paragraph()
            
            # Opening Ceremony
            doc.add_heading("Opening Ceremony", level=2)
            if self.data['opening_ceremony']['prayer_intentions']:
                doc.add_paragraph(f"Prayer Intentions: {self.data['opening_ceremony']['prayer_intentions']}")
            if self.data['opening_ceremony']['opening_prayer']:
                doc.add_paragraph(f"Opening Prayer: {self.data['opening_ceremony']['opening_prayer']}")
            if self.data['opening_ceremony']['led_by']:
                doc.add_paragraph(f"Led by: {self.data['opening_ceremony']['led_by']}")
            if self.data['opening_ceremony']['pledge_of_allegiance']:
                doc.add_paragraph(f"Pledge of Allegiance led by: {self.data['opening_ceremony']['pledge_of_allegiance']}")
            doc.add_paragraph()
            
            # Roll Call
            doc.add_heading("Roll Call", level=2)
            present = []
            absent = []
            excused = []
            
            for office, info in self.data['roll_call'].items():
                attendance = info['attendance']
                name = info['name']
                if attendance == "Present":
                    present.append(f"{office} - {name}")
                elif attendance == "Absent":
                    absent.append(f"{office} - {name}")
                elif attendance == "Excused":
                    excused.append(f"{office} - {name}")
            
            if present:
                doc.add_paragraph("Present:", style='List Bullet')
                for member in present:
                    doc.add_paragraph(member, style='List Bullet 2')
            
            if excused:
                doc.add_paragraph("Excused:", style='List Bullet')
                for member in excused:
                    doc.add_paragraph(member, style='List Bullet 2')
                    
            if absent:
                doc.add_paragraph("Absent:", style='List Bullet')
                for member in absent:
                    doc.add_paragraph(member, style='List Bullet 2')
            doc.add_paragraph()
            
            # Faithful Friar
            if self.data['meeting_notes']['Faithful Friar']:
                doc.add_heading("Faithful Friar", level=2)
                doc.add_paragraph(self.data['meeting_notes']['Faithful Friar'])
                doc.add_paragraph()
            
            # Reading of Minutes
            doc.add_heading("Reading of Minutes", level=2)
            minutes_data = self.data['meeting_notes']['Reading of Minutes']
            if minutes_data['corrections']:
                doc.add_paragraph(f"Corrections: {minutes_data['corrections']}")
            if minutes_data['motion_to_approve']:
                doc.add_paragraph(f"Motion to approve by: {minutes_data['motion_to_approve']}")
            if minutes_data['seconded_by']:
                doc.add_paragraph(f"Seconded by: {minutes_data['seconded_by']}")
            doc.add_paragraph(f"Minutes approved: {minutes_data['approved']}")
            doc.add_paragraph()
            
            # Bills and Communications
            if self.data['meeting_notes']['Bills and Communications']:
                doc.add_heading("Bills and Communications", level=2)
                doc.add_paragraph(self.data['meeting_notes']['Bills and Communications'])
                doc.add_paragraph()
            
            # Faithful Comptroller
            if self.data['meeting_notes']['Faithful Comptroller']:
                doc.add_heading("Faithful Comptroller", level=2)
                doc.add_paragraph(self.data['meeting_notes']['Faithful Comptroller'])
                doc.add_paragraph()
            
            # Purser's Report
            if self.data['meeting_notes']["Purser's Report"]:
                doc.add_heading("Purser's Report", level=2)
                doc.add_paragraph(self.data['meeting_notes']["Purser's Report"])
                doc.add_paragraph()
            
            # Standing Committees
            if self.data['meeting_notes']['Standing Committees']['Color Corps']:
                doc.add_heading("Standing Committees", level=2)
                doc.add_heading("Color Corps", level=3)
                doc.add_paragraph(self.data['meeting_notes']['Standing Committees']['Color Corps'])
                doc.add_paragraph()
            
            # Reading of Applications
            if self.data['meeting_notes']['Reading of Applications']:
                doc.add_heading("Reading of Applications", level=2)
                doc.add_paragraph(self.data['meeting_notes']['Reading of Applications'])
                doc.add_paragraph()
            
            # Unfinished Business
            if self.data['meeting_notes']['Unfinished Business']:
                doc.add_heading("Unfinished Business", level=2)
                doc.add_paragraph(self.data['meeting_notes']['Unfinished Business'])
                doc.add_paragraph()
            
            # New Business
            if self.data['meeting_notes']['New Business']:
                doc.add_heading("New Business", level=2)
                doc.add_paragraph(self.data['meeting_notes']['New Business'])
                doc.add_paragraph()
            
            # Trustee's Report
            if self.data['meeting_notes']["Trustee's Report"]:
                doc.add_heading("Trustee's Report", level=2)
                doc.add_paragraph(self.data['meeting_notes']["Trustee's Report"])
                doc.add_paragraph()
            
            # Report of the Councils
            has_council_reports = any(self.data['meeting_notes']['Report of the Councils'].values())
            if has_council_reports:
                doc.add_heading("Report of the Councils", level=2)
                for council, report in self.data['meeting_notes']['Report of the Councils'].items():
                    if report:
                        doc.add_heading(council, level=3)
                        doc.add_paragraph(report)
                doc.add_paragraph()
            
            # Closing
            doc.add_heading("Closing", level=2)
            if self.data['closing_notes']['time']:
                doc.add_paragraph(f"Meeting adjourned at: {self.data['closing_notes']['time']}")
            if self.data['closing_notes']['Next Officers Meeting']:
                doc.add_paragraph(f"Next Officers Meeting: {self.data['closing_notes']['Next Officers Meeting']}")
            if self.data['closing_notes']['Next Business Meeting']:
                doc.add_paragraph(f"Next Business Meeting: {self.data['closing_notes']['Next Business Meeting']}")
            
            # Save document
            doc.save(filename)
            messagebox.showinfo("Success", f"Minutes exported to:\n{filename}")
            
        except Exception as e:
            messagebox.showerror("Export Error", f"An error occurred:\n{str(e)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = KofCMinutesApp(root)
    root.mainloop()